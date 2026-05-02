"""
统一文本抽取：PDF（PyMuPDF 优先，失败则 pypdf）、Word .docx（python-docx）。
不含 OCR；扫描版 PDF 可能仍为空。
"""
from __future__ import annotations

from io import BytesIO
from typing import Optional


def extract_pdf_bytes(data: bytes) -> str:
    """从 PDF 二进制抽取纯文本：PyMuPDF 优先，正文过少时用 pypdf 补足。"""
    t1 = (_extract_pdf_pymupdf(data) or "").strip()
    t2 = (_extract_pdf_pypdf(data) or "").strip()
    if len(t1) >= 30:
        return t1
    if len(t2) >= 30:
        return t2
    merged = "\n\n".join(x for x in (t1, t2) if x).strip()
    return merged


def _extract_pdf_pymupdf(data: bytes) -> Optional[str]:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return None
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        parts = []
        for i in range(len(doc)):
            parts.append(doc[i].get_text() or "")
        doc.close()
        return "\n\n".join(parts)
    except Exception:
        return None


def _extract_pdf_pypdf(data: bytes) -> Optional[str]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return None
    try:
        reader = PdfReader(BytesIO(data))
        parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n\n".join(parts)
    except Exception:
        return None


def extract_docx_bytes(data: bytes) -> str:
    """抽取 Word .docx 正文（段落 + 简单表格单元格）。"""
    import io

    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            lines.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                c = (cell.text or "").strip()
                if c:
                    lines.append(c)
    return "\n".join(lines)
